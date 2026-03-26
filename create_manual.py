"""Generate user manual for Screenshot Monitor (Electron version) as a Word document."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Meiryo UI"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

# -- Title --
title = doc.add_heading("Screenshot Monitor 操作マニュアル", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_ver = doc.add_paragraph("Version 1.0.0")
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# ========================================
# 1. Overview
# ========================================
doc.add_heading("1. 概要", level=1)
doc.add_paragraph(
    "Screenshot Monitor は、画面上の指定した領域を自動で監視し、"
    "変化を検知するとスクリーンショットを自動撮影するデスクトップアプリです。\n"
    "撮影した画像は一覧で確認・選択でき、ワンクリックでPDFに書き出すことができます。"
)

doc.add_heading("主な機能", level=2)
features = [
    "画面上の任意の領域をドラッグで指定して監視",
    "マルチモニター対応（メイン画面・サブ画面）",
    "画面の変化を自動検知してスクリーンショットを撮影",
    "撮影時に領域が白くフラッシュして撮影を通知",
    "監視中に四隅のL字ハンドルで領域をリサイズ可能",
    "一時停止 / 再開が可能",
    "撮影画像の一覧表示・クリックで選択 / 非選択の切り替え",
    "ドラッグで複数画像を一括選択 / 一括非選択",
    "拡大ビューアで画像を確認（キーボード操作対応）",
    "選択した画像をPDFに一括書き出し",
    "既存フォルダの画像を読み込んでPDF化",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

# ========================================
# 2. Install
# ========================================
doc.add_heading("2. インストール方法", level=1)
doc.add_paragraph(
    "配布された「Screenshot Monitor Setup 1.0.0.exe」をダブルクリックしてください。\n"
    "自動でインストールが完了し、アプリが起動します。"
)
p_note = doc.add_paragraph()
p_note.add_run("注意: ").bold = True
p_note.add_run(
    "初回ダウンロード時に「WindowsによってPCが保護されました」と表示される場合があります。\n"
    "「詳細情報」をクリック →「実行」をクリックしてください。"
)

doc.add_paragraph("")
doc.add_paragraph(
    "インストール後はスタートメニューから「Screenshot Monitor」で起動できます。\n"
    "アンインストールは「設定」→「アプリ」→「Screenshot Monitor」から行えます。"
)

# ========================================
# 3. Initial Screen
# ========================================
doc.add_heading("3. 初期画面の説明", level=1)
doc.add_paragraph(
    "アプリを起動すると、以下のボタンと設定項目が表示されます。"
)

table = doc.add_table(rows=4, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = table.rows[0].cells
cells[0].text = "ボタン"
cells[1].text = "説明"
data = [
    ("キャプチャ領域を選択", "監視する画面範囲をドラッグで指定します"),
    ("監視スタート（赤いボタン）", "選択した領域の監視を開始します（領域選択後に有効化）"),
    ("既存フォルダを参照", "過去に撮影した画像フォルダを開いてPDF書き出しができます"),
]
for i, (btn, desc) in enumerate(data, 1):
    cells = table.rows[i].cells
    cells[0].text = btn
    cells[1].text = desc

doc.add_paragraph("")

doc.add_heading("設定項目", level=2)
table2 = doc.add_table(rows=4, cols=3)
table2.style = "Light Grid Accent 1"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
header = table2.rows[0].cells
header[0].text = "項目"
header[1].text = "デフォルト"
header[2].text = "説明"
settings = [
    ("チェック間隔", "3秒",
     "画面の変化を確認する頻度です。短いほど細かく検知します。"),
    ("変化の割合", "15%",
     "監視領域の何%のピクセルが変化したら撮影するかの基準です。\n"
     "値を大きくすると、小さな変化（一部だけの変更）を無視できます。\n"
     "例: 30%にすると領域の30%以上が変わらない限り撮影しません。"),
    ("色の感度", "30",
     "1ピクセルの色の変化がこの値（0-255）以上の場合に「変化した」とみなします。\n"
     "値を大きくすると微細な色の変化を無視できます。"),
]
for i, (name, default, desc) in enumerate(settings, 1):
    cells = table2.rows[i].cells
    cells[0].text = name
    cells[1].text = default
    cells[2].text = desc

# ========================================
# 4. Operation Steps
# ========================================
doc.add_heading("4. 操作手順", level=1)

# Step 1
doc.add_heading("Step 1: キャプチャ領域を選択", level=2)
doc.add_paragraph(
    "「キャプチャ領域を選択」ボタンをクリックすると、画面全体が暗くなります。\n"
    "監視したい範囲をマウスでドラッグして選択してください。\n"
    "マルチモニター環境では、メイン画面・サブ画面のどちらでも選択可能です。\n"
    "選択後、赤い枠線で領域が表示されます。"
)
doc.add_paragraph("")
p = doc.add_paragraph()
p.add_run("ポイント: ").bold = True
p.add_run(
    "赤枠の四隅にL字型の白いハンドルが表示されます。"
    "このハンドルをドラッグすることで、領域のサイズや位置を微調整できます。"
    "この段階では、赤枠はメインウィンドウの背面に表示されるため、"
    "ウィンドウの操作を妨げません。"
)

# Step 2
doc.add_heading("Step 2: 監視スタート", level=2)
doc.add_paragraph(
    "設定を確認し、赤い「監視スタート」ボタンをクリックします。\n"
    "メインウィンドウは非表示になり、赤枠の右上にコントロールパネルが表示されます。"
)

doc.add_paragraph("")
doc.add_heading("コントロールパネルの表示内容", level=3)
table3 = doc.add_table(rows=4, cols=2)
table3.style = "Light Grid Accent 1"
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
header = table3.rows[0].cells
header[0].text = "表示"
header[1].text = "説明"
ctrl_data = [
    ("● 録画監視中", "赤い丸が点滅して監視中であることを示します"),
    ("停止ボタン",
     "監視を一時停止します。ボタンが「再開」に変わります。\n"
     "一時停止中は赤い丸が黄色に変わります。"),
    ("中止ボタン", "監視を終了し、撮影結果の一覧画面に移動します"),
]
for i, (label, desc) in enumerate(ctrl_data, 1):
    cells = table3.rows[i].cells
    cells[0].text = label
    cells[1].text = desc

doc.add_paragraph("")
p2 = doc.add_paragraph()
p2.add_run("撮影の合図: ").bold = True
p2.add_run("スクリーンショットが撮影されると、監視領域が一瞬白くフラッシュします。")

doc.add_paragraph("")
p3 = doc.add_paragraph()
p3.add_run("監視中のリサイズ: ").bold = True
p3.add_run(
    "監視中でも四隅のハンドルをドラッグして領域を変更できます。"
    "変更後は自動的にスクリーンショットが撮影されます。"
)

# Step 3
doc.add_heading("Step 3: 結果の確認", level=2)
doc.add_paragraph(
    "「中止」ボタンで監視を終了すると、撮影したスクリーンショットの一覧が表示されます。\n"
    "各画像の左上にチェックマーク（✓）が表示されます。"
)

doc.add_heading("画像の選択 / 非選択", level=3)
select_ops = [
    ("クリック", "画像を1つずつ選択 / 非選択を切り替えます"),
    ("左ドラッグ（非選択画像から開始）",
     "青い枠で囲んだ画像をまとめて選択します"),
    ("左ドラッグ（選択済み画像から開始）",
     "赤い枠で囲んだ画像をまとめて非選択にします"),
    ("全て選択ボタン", "全画像を選択状態にします"),
    ("全て解除ボタン", "全画像を非選択状態にします"),
]
table_sel = doc.add_table(rows=len(select_ops) + 1, cols=2)
table_sel.style = "Light Grid Accent 1"
table_sel.alignment = WD_TABLE_ALIGNMENT.CENTER
table_sel.rows[0].cells[0].text = "操作"
table_sel.rows[0].cells[1].text = "動作"
for i, (op, desc) in enumerate(select_ops, 1):
    table_sel.rows[i].cells[0].text = op
    table_sel.rows[i].cells[1].text = desc

doc.add_paragraph("")

doc.add_heading("拡大ビューア", level=3)
doc.add_paragraph(
    "画像にマウスを合わせると右上に虫眼鏡アイコンが表示されます。\n"
    "クリックすると拡大表示されます。"
)

viewer_ops = [
    ("← → キー", "前後の画像に移動"),
    ("スペースキー", "選択 / 非選択を切り替え"),
    ("左上のチェックマーク", "クリックで選択 / 非選択を切り替え"),
    ("Escキー / 背景クリック", "ビューアを閉じる"),
]
table_viewer = doc.add_table(rows=len(viewer_ops) + 1, cols=2)
table_viewer.style = "Light Grid Accent 1"
table_viewer.alignment = WD_TABLE_ALIGNMENT.CENTER
table_viewer.rows[0].cells[0].text = "操作"
table_viewer.rows[0].cells[1].text = "動作"
for i, (op, desc) in enumerate(viewer_ops, 1):
    table_viewer.rows[i].cells[0].text = op
    table_viewer.rows[i].cells[1].text = desc

doc.add_paragraph("")

doc.add_heading("結果画面のボタン", level=3)
table4 = doc.add_table(rows=6, cols=2)
table4.style = "Light Grid Accent 1"
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
table4.rows[0].cells[0].text = "ボタン"
table4.rows[0].cells[1].text = "説明"
result_data = [
    ("新規録画", "新しいセッションを開始します（初期画面に戻る）"),
    ("既存フォルダを参照", "別のフォルダの画像を読み込みます"),
    ("PDFに書き出す", "選択中の画像をPDFファイルとして保存します"),
    ("全て選択", "全画像を選択状態にします"),
    ("全て解除", "全画像を非選択状態にします"),
]
for i, (btn, desc) in enumerate(result_data, 1):
    table4.rows[i].cells[0].text = btn
    table4.rows[i].cells[1].text = desc

# Step 4
doc.add_heading("Step 4: PDF書き出し", level=2)
doc.add_paragraph(
    "「PDFに書き出す」ボタンをクリックすると、保存先を選択するダイアログが表示されます。\n"
    "ファイル名と保存場所を指定して保存してください。\n"
    "選択中（チェックマーク付き）の画像のみがPDFに含まれます。\n"
    "各ページに1枚ずつ、画像がページいっぱいに配置されます。"
)

# ========================================
# 5. Folder Structure
# ========================================
doc.add_heading("5. フォルダ構成", level=1)
doc.add_paragraph(
    "スクリーンショットはアプリと同じ場所にある「Screenshots」フォルダに保存されます。\n"
    "セッション（録画）ごとにフォルダが自動作成されます。"
)

lines = [
    "Screenshot Monitor/",
    "    Screenshots/",
    "        session_20260326_100000/",
    "            screenshot_20260326_100003_123.png",
    "            screenshot_20260326_100006_456.png",
    "        session_20260326_140000/",
    "            ...",
]
for line in lines:
    p = doc.add_paragraph()
    p.style = "No Spacing"
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(line)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

# ========================================
# 6. Tips & FAQ
# ========================================
doc.add_heading("6. よくある質問・Tips", level=1)

doc.add_heading("画面の一部だけ変わった場合にスクショを撮りたくない", level=2)
doc.add_paragraph(
    "「変化の割合」を大きくしてください。"
    "例えば30%にすると、領域全体の30%以上が変わらない限りスクリーンショットは撮影されません。"
)

doc.add_heading("スクリーンショットの撮影頻度を変えたい", level=2)
doc.add_paragraph(
    "「チェック間隔」を変更してください。"
    "1秒にすると毎秒チェック、10秒にすると10秒ごとにチェックします。"
)

doc.add_heading("微細な色の変化を無視したい", level=2)
doc.add_paragraph(
    "「色の感度」を大きくしてください。"
    "例えば50にすると、各ピクセルの色差が50以上でないと変化として認識しません。"
)

doc.add_heading("過去のスクリーンショットからPDFを作りたい", level=2)
doc.add_paragraph(
    "「既存フォルダを参照」から対象のセッションフォルダを選択してください。"
    "フォルダ内の画像が一覧表示され、選択してPDFに書き出すことができます。"
)

doc.add_heading("複数の画像をまとめて非選択にしたい", level=2)
doc.add_paragraph(
    "選択済みの画像の上から左ドラッグを開始すると、"
    "赤い枠が表示され、囲んだ画像をまとめて非選択にできます。"
)

doc.add_heading("アンインストールしたい", level=2)
doc.add_paragraph(
    "「設定」→「アプリ」→「Screenshot Monitor」→「アンインストール」で削除できます。\n"
    "Screenshots フォルダ内の画像は自動削除されません。必要に応じて手動で削除してください。"
)

# ========================================
# 7. Download
# ========================================
doc.add_heading("7. ダウンロード", level=1)
doc.add_paragraph(
    "最新版は以下のURLからダウンロードできます。"
)
doc.add_paragraph(
    "https://github.com/gegege1566/screenshot-monitor/releases"
)

# -- Save --
output_path = r"C:\Users\K22030126\gegege2030\screenshot-monitor-electron\Screenshot_Monitor_Manual.docx"
doc.save(output_path)
print(f"Manual saved: {output_path}")
